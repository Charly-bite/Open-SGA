#!/usr/bin/env python3
"""
HDS Database Update Script
===========================
Takes the discrepancies found by audit_hds.py and updates the
Unified_GHS_Database.csv to match the HDS source documents.

Creates a backup before making changes and generates a detailed
change log.
"""

import csv
import os
import re
import sys
import shutil
import unicodedata
from pathlib import Path
from datetime import datetime

try:
    import fitz

    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

try:
    import docx

    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# ============================================================
# CONFIGURATION
# ============================================================
BASE_DIR = Path(r"c:\Users\QB_DESARROLLO\Desktop\SGA_dev")
MATRIZ_CSV = BASE_DIR / "data" / "MATRIZ - HDS Y FT (ACTUALIZACION)(Hoja2).csv"
DATABASE_CSV = BASE_DIR / "original_data" / "Unified_GHS_Database.csv"
HDS_FOLDER = Path(r"C:\Users\QB_DESARROLLO\Desktop\Nueva carpeta")
CHANGELOG = BASE_DIR / "scripts" / "update_changelog.txt"
PREVIEW_FILE = BASE_DIR / "scripts" / "update_preview.txt"

# Database columns
DB_HEADERS = [
    "Codigo interno",
    "Chemical Name",
    "CAS Number",
    "Signal Word",
    "Emergencia",
    "H-Statements",
    "Consejos (Frases P)",
    "Source_File",
    "Exclamación",
    "Corrosión",
    "Peligro para la salud",
    "Calavera",
    "Llama",
    "Ambiente",
    "Llama sobre círculo",
    "Bomba explotando",
    "Cilindro de gas",
]


# ============================================================
# TEXT UTILITIES (same as audit_hds.py)
# ============================================================
def normalize_text(text):
    if not text:
        return ""
    nfkd = unicodedata.normalize("NFKD", text)
    text = "".join(c for c in nfkd if not unicodedata.combining(c))
    text = text.lower().strip()
    text = re.sub(r"\s+", " ", text)
    return text


def normalize_for_matching(name):
    n = normalize_text(name)
    n = re.sub(r"\bs/?p\b", "", n)
    n = re.sub(r"\bqb\b", "", n)
    n = re.sub(r"[^\w\s]", " ", n)
    n = re.sub(r"\s+", " ", n).strip()
    return n


def normalize_filename(filepath):
    name = Path(filepath).stem
    name = re.sub(r"^HDS\s*[-–—]\s*", "", name, flags=re.IGNORECASE)
    name = re.sub(r"^HDS\s+", "", name, flags=re.IGNORECASE)
    name = re.sub(r"^FORMATO\s+HDS\s+", "", name, flags=re.IGNORECASE)
    name = re.sub(r"\s+\d+$", "", name)
    name = re.sub(r"\s+[Aa]$", "", name)
    return normalize_for_matching(name)


# ============================================================
# TEXT EXTRACTION
# ============================================================
def extract_text_from_pdf(filepath):
    if not HAS_PYMUPDF:
        return ""
    try:
        doc = fitz.open(str(filepath))
        text_parts = []
        for page in doc:
            text_parts.append(page.get_text())
        doc.close()
        return "\n".join(text_parts)
    except Exception:
        return ""


def extract_text_from_docx(filepath):
    if not HAS_DOCX:
        return ""
    try:
        doc = docx.Document(str(filepath))
        text_parts = []
        for para in doc.paragraphs:
            text_parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_parts.append(cell.text)
        return "\n".join(text_parts)
    except Exception:
        return ""


def extract_text(filepath):
    ext = Path(filepath).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(filepath)
    elif ext == ".docx":
        return extract_text_from_docx(filepath)
    return ""


# ============================================================
# SAFETY DATA EXTRACTION
# ============================================================
def extract_h_statements(text):
    matches = re.findall(r"\bH\s*(\d{3})\b", text, re.IGNORECASE)
    return sorted(set(f"H{m}" for m in matches))


def extract_p_statements(text):
    matches = re.findall(r"\bP\s*(\d{3})\b", text, re.IGNORECASE)
    return sorted(set(f"P{m}" for m in matches))


def extract_signal_word(text):
    text_upper = text.upper()
    if re.search(r"\bPELIGRO\b", text_upper):
        return "PELIGRO"
    if re.search(r"\bDANGER\b", text_upper):
        return "PELIGRO"
    if re.search(r"\bATENCI[OÓ]N\b", text_upper):
        return "ATENCION"
    if re.search(r"\bWARNING\b", text_upper):
        return "ATENCION"
    return "No Aplicable"


def extract_cas_number(text):
    matches = re.findall(r"\b(\d{2,7}-\d{2}-\d)\b", text)
    return sorted(set(matches))


# ============================================================
# PICTOGRAM DETERMINATION FROM H-STATEMENTS
# ============================================================
# Maps H-statement ranges to GHS pictograms
H_TO_PICTOGRAM = {
    # Exclamación (GHS07)
    "Exclamación": {
        "H302",
        "H303",
        "H312",
        "H313",
        "H315",
        "H316",
        "H317",
        "H319",
        "H320",
        "H332",
        "H333",
        "H335",
        "H336",
        "H402",
        "H412",
    },
    # Corrosión (GHS05)
    "Corrosión": {
        "H290",
        "H314",
        "H318",
    },
    # Peligro para la salud (GHS08)
    "Peligro para la salud": {
        "H304",
        "H334",
        "H340",
        "H341",
        "H350",
        "H351",
        "H360",
        "H361",
        "H370",
        "H371",
        "H372",
        "H373",
    },
    # Calavera (GHS06)
    "Calavera": {
        "H300",
        "H301",
        "H310",
        "H311",
        "H330",
        "H331",
    },
    # Llama (GHS02)
    "Llama": {
        "H220",
        "H221",
        "H222",
        "H223",
        "H224",
        "H225",
        "H226",
        "H227",
        "H228",
        "H241",
        "H242",
        "H250",
        "H251",
        "H252",
        "H260",
        "H261",
    },
    # Ambiente (GHS09)
    "Ambiente": {
        "H400",
        "H401",
        "H410",
        "H411",
    },
    # Llama sobre círculo (GHS03)
    "Llama sobre círculo": {
        "H270",
        "H271",
        "H272",
    },
    # Bomba explotando (GHS01)
    "Bomba explotando": {
        "H200",
        "H201",
        "H202",
        "H203",
        "H204",
        "H205",
        "H240",
        "H241",
    },
    # Cilindro de gas (GHS04)
    "Cilindro de gas": {
        "H280",
        "H281",
        "H282",
    },
}


def determine_pictograms(h_codes):
    """Determine which pictograms should be active based on H-statements."""
    pictos = {}
    for picto_name, h_set in H_TO_PICTOGRAM.items():
        active = bool(set(h_codes) & h_set)
        pictos[picto_name] = "X" if active else ""
    return pictos


# ============================================================
# FILE MATCHING (same as audit_hds.py)
# ============================================================
def build_hds_index(hds_folder):
    index = {}
    file_list = []
    for root, dirs, files in os.walk(hds_folder):
        for f in files:
            ext = Path(f).suffix.lower()
            if ext in (".pdf", ".docx"):
                full_path = os.path.join(root, f)
                norm_name = normalize_filename(f)
                file_list.append((norm_name, full_path, f))
                if norm_name not in index:
                    index[norm_name] = []
                index[norm_name].append(full_path)
    return index, file_list


def find_hds_file(product_name, hds_index, hds_file_list):
    norm_product = normalize_for_matching(product_name)
    if not norm_product:
        return None
    if norm_product in hds_index:
        return hds_index[norm_product][0]
    best_match = None
    best_score = 0
    for norm_name, filepath, original_name in hds_file_list:
        if not norm_name:
            continue
        if norm_product in norm_name or norm_name in norm_product:
            score = min(len(norm_product), len(norm_name)) / max(
                len(norm_product), len(norm_name)
            )
            if score > best_score:
                best_score = score
                best_match = filepath
        product_words = set(norm_product.split())
        file_words = set(norm_name.split())
        stop_words = {
            "de",
            "la",
            "el",
            "en",
            "y",
            "a",
            "con",
            "para",
            "por",
            "del",
            "las",
            "los",
            "un",
            "una",
        }
        product_words -= stop_words
        file_words -= stop_words
        if product_words and file_words:
            common = product_words & file_words
            if common:
                score = len(common) / max(len(product_words), len(file_words))
                if score > best_score and score >= 0.5:
                    best_score = score
                    best_match = filepath
    if best_score >= 0.4:
        return best_match
    return None


# ============================================================
# LOAD CSV DATA
# ============================================================
def load_matriz(filepath):
    products = []
    for enc in ["utf-8-sig", "utf-8", "latin-1", "cp1252"]:
        try:
            with open(filepath, "r", encoding=enc) as f:
                reader = csv.DictReader(f)
                for row in reader:
                    etiqueta = (row.get("ETIQUETA") or "").strip()
                    if etiqueta == "REVISADO":
                        products.append(
                            {
                                "codigo": (row.get("CODIGO") or "").strip(),
                                "nombre": (row.get("NOMBRE") or "").strip(),
                            }
                        )
            return products
        except (UnicodeDecodeError, UnicodeError):
            continue
    return products


def load_database_raw(filepath):
    """Load database as raw rows for modification."""
    for enc in ["utf-8-sig", "utf-8", "latin-1", "cp1252"]:
        try:
            with open(filepath, "r", encoding=enc, newline="") as f:
                content = f.read()
            return content, enc
        except (UnicodeDecodeError, UnicodeError):
            continue
    return None, None


def parse_database(content):
    """Parse database content into rows, handling multiline fields."""
    reader = csv.reader(content.splitlines())
    headers = None
    rows = []

    for line_parts in reader:
        if headers is None:
            headers = [h.strip() for h in line_parts]
            continue

        # A new row starts when the first column has content
        if line_parts and line_parts[0].strip():
            rows.append(line_parts)
        elif rows:
            # Continuation line - merge into last row
            last = rows[-1]
            for i, part in enumerate(line_parts):
                if i < len(last):
                    last[i] += " " + part
                else:
                    last.append(part)

    return headers, rows


# ============================================================
# MAIN UPDATE LOGIC
# ============================================================
def run_update(dry_run=True):
    """Run the database update. If dry_run=True, only preview changes."""

    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    mode = "PREVIEW" if dry_run else "UPDATE"

    print(f"{'=' * 70}")
    print(f"  HDS DATABASE {mode}")
    print(f"  {timestamp}")
    print(f"{'=' * 70}")

    # 1. Load all data sources
    print("\n[1/4] Loading data...")
    products = load_matriz(MATRIZ_CSV)
    print(f"  → {len(products)} REVISADO products")

    content, encoding = load_database_raw(DATABASE_CSV)
    if not content:
        print("ERROR: Could not read database file")
        return

    headers, rows = parse_database(content)
    print(f"  → {len(rows)} database rows (encoding: {encoding})")

    # Build lookup: codigo -> row index(es)
    codigo_col = headers.index("Codigo interno") if "Codigo interno" in headers else 0
    db_lookup = {}
    for idx, row in enumerate(rows):
        if len(row) > codigo_col:
            code = row[codigo_col].strip()
            if code:
                if code not in db_lookup:
                    db_lookup[code] = []
                db_lookup[code].append(idx)

    # Column indices
    col_map = {h: i for i, h in enumerate(headers)}

    hds_index, hds_file_list = build_hds_index(HDS_FOLDER)
    print(f"  → {len(hds_file_list)} HDS files indexed")

    # 2. Find changes needed
    print("\n[2/4] Analyzing discrepancies...")

    changes = (
        []
    )  # List of (row_idx, field, old_value, new_value, product_name, codigo, hds_file)

    for product in products:
        codigo = product["codigo"]
        nombre = product["nombre"]

        if codigo not in db_lookup:
            continue

        # Find HDS file
        hds_file = find_hds_file(nombre, hds_index, hds_file_list)
        if not hds_file:
            db_name_idx = col_map.get("Chemical Name", 1)
            for ridx in db_lookup[codigo]:
                if len(rows[ridx]) > db_name_idx:
                    db_name = rows[ridx][db_name_idx].strip()
                    hds_file = find_hds_file(db_name, hds_index, hds_file_list)
                    if hds_file:
                        break

        if not hds_file:
            continue

        # Extract HDS data
        text = extract_text(hds_file)
        if not text:
            continue

        hds_h = extract_h_statements(text)
        hds_signal = extract_signal_word(text)
        hds_p = extract_p_statements(text)
        hds_cas = extract_cas_number(text)

        # Compare against each DB row for this codigo
        for ridx in db_lookup[codigo]:
            row = rows[ridx]

            # Check source - prefer updating Master rows
            source_idx = col_map.get("Source_File", -1)
            source = (
                row[source_idx].strip()
                if source_idx >= 0 and len(row) > source_idx
                else ""
            )

            # --- Signal Word ---
            sw_idx = col_map.get("Signal Word", -1)
            if sw_idx >= 0 and len(row) > sw_idx:
                db_signal = row[sw_idx].strip()
                db_signal_norm = normalize_text(db_signal)
                hds_signal_norm = normalize_text(hds_signal)

                if (
                    db_signal_norm != hds_signal_norm
                    and hds_signal_norm != "no aplicable"
                    and db_signal not in ("", "-")
                ):
                    changes.append(
                        {
                            "row_idx": ridx,
                            "col_idx": sw_idx,
                            "field": "Signal Word",
                            "old": db_signal,
                            "new": hds_signal,
                            "nombre": nombre,
                            "codigo": codigo,
                            "hds_file": Path(hds_file).name,
                            "source": source,
                        }
                    )
                elif db_signal in ("", "-") and hds_signal_norm != "no aplicable":
                    changes.append(
                        {
                            "row_idx": ridx,
                            "col_idx": sw_idx,
                            "field": "Signal Word",
                            "old": db_signal or "(empty)",
                            "new": hds_signal,
                            "nombre": nombre,
                            "codigo": codigo,
                            "hds_file": Path(hds_file).name,
                            "source": source,
                        }
                    )

            # --- H-Statements ---
            h_idx = col_map.get("H-Statements", -1)
            if h_idx >= 0 and len(row) > h_idx and hds_h:
                db_h_raw = row[h_idx].strip()
                db_h_codes = set(extract_h_statements(db_h_raw))
                hds_h_set = set(hds_h)

                missing_in_db = hds_h_set - db_h_codes

                if missing_in_db:
                    # Build new H-statements string: merge existing + missing
                    all_h = sorted(db_h_codes | hds_h_set)
                    new_h_str = ", ".join(all_h)

                    changes.append(
                        {
                            "row_idx": ridx,
                            "col_idx": h_idx,
                            "field": "H-Statements",
                            "old": db_h_raw[:80]
                            + ("..." if len(db_h_raw) > 80 else ""),
                            "new": new_h_str,
                            "nombre": nombre,
                            "codigo": codigo,
                            "hds_file": Path(hds_file).name,
                            "source": source,
                            "detail": f"Added: {', '.join(sorted(missing_in_db))}",
                        }
                    )

            # --- Pictograms (recalculate based on final H-statements) ---
            if h_idx >= 0 and len(row) > h_idx and hds_h:
                all_h = sorted(set(extract_h_statements(row[h_idx])) | set(hds_h))
                new_pictos = determine_pictograms(all_h)

                picto_cols = [
                    "Exclamación",
                    "Corrosión",
                    "Peligro para la salud",
                    "Calavera",
                    "Llama",
                    "Ambiente",
                    "Llama sobre círculo",
                    "Bomba explotando",
                    "Cilindro de gas",
                ]

                for pcol in picto_cols:
                    pidx = col_map.get(pcol, -1)
                    if pidx < 0:
                        continue

                    old_val = row[pidx].strip().upper() if len(row) > pidx else ""
                    old_active = old_val in ("X", "1")
                    new_active = new_pictos.get(pcol, "") == "X"

                    if old_active != new_active:
                        changes.append(
                            {
                                "row_idx": ridx,
                                "col_idx": pidx,
                                "field": f"Pictogram: {pcol}",
                                "old": "X" if old_active else "(empty)",
                                "new": "X" if new_active else "(empty)",
                                "nombre": nombre,
                                "codigo": codigo,
                                "hds_file": Path(hds_file).name,
                                "source": source,
                            }
                        )

    # 3. Report changes
    print(
        f"\n  → Found {len(changes)} changes needed across {len(set(c['codigo'] for c in changes))} products"
    )

    # Group by product
    by_product = {}
    for c in changes:
        key = c["codigo"]
        if key not in by_product:
            by_product[key] = []
        by_product[key].append(c)

    # Write preview/changelog
    output_file = PREVIEW_FILE if dry_run else CHANGELOG
    with open(output_file, "w", encoding="utf-8") as f:
        f.write(f"{'=' * 80}\n")
        f.write(f"  HDS DATABASE {mode}\n")
        f.write(f"  {timestamp}\n")
        f.write(f"{'=' * 80}\n\n")
        f.write(f"Total changes: {len(changes)}\n")
        f.write(f"Products affected: {len(by_product)}\n\n")

        for codigo, product_changes in sorted(by_product.items()):
            nombre = product_changes[0]["nombre"]
            hds = product_changes[0]["hds_file"]
            f.write(f"{'─' * 60}\n")
            f.write(f"PRODUCT: {nombre} [{codigo}]\n")
            f.write(f"HDS Source: {hds}\n")

            for c in product_changes:
                detail = c.get("detail", "")
                f.write(f"  {c['field']}:\n")
                f.write(f"    OLD: {c['old']}\n")
                f.write(f"    NEW: {c['new']}\n")
                if detail:
                    f.write(f"    ({detail})\n")
            f.write("\n")

        f.write(f"\n{'=' * 80}\n")
        f.write(f"  END OF {mode}\n")
        f.write(f"{'=' * 80}\n")

    print(f"\n  {mode} saved to: {output_file}")

    # 4. Apply changes if not dry_run
    if not dry_run:
        print("\n[3/4] Creating backup...")
        backup_path = DATABASE_CSV.with_suffix(
            f'.csv.backup_{datetime.now().strftime("%Y%m%d_%H%M%S")}'
        )
        shutil.copy2(DATABASE_CSV, backup_path)
        print(f"  → Backup saved to: {backup_path}")

        print("\n[4/4] Applying changes...")

        # Apply all changes to rows
        for c in changes:
            ridx = c["row_idx"]
            cidx = c["col_idx"]
            # Ensure row is long enough
            while len(rows[ridx]) <= cidx:
                rows[ridx].append("")
            rows[ridx][cidx] = c["new"]

        # Write updated CSV
        with open(DATABASE_CSV, "w", encoding=encoding, newline="") as f:
            writer = csv.writer(f)
            writer.writerow(headers)
            for row in rows:
                writer.writerow(row)

        print(f"  → Database updated: {DATABASE_CSV}")
        print(f"  → {len(changes)} changes applied across {len(by_product)} products")
    else:
        print(f"\n  This is a DRY RUN. No changes were made.")
        print(f"  Review the preview file, then run with --apply to apply changes:")
        print(f"  python scripts/update_hds.py --apply")

    print(f"\n{'=' * 70}")
    print(f"  {mode} COMPLETE")
    print(f"{'=' * 70}")


if __name__ == "__main__":
    dry_run = "--apply" not in sys.argv
    run_update(dry_run=dry_run)
