#!/usr/bin/env python3
"""
HDS Database Audit Script
=========================
Audits the SGA Product Database entries against the original HDS (Safety Data Sheet)
source documents.

Workflow:
1. Read MATRIZ CSV → filter products where ETIQUETA = "REVISADO"
2. Look up each product in Unified_GHS_Database.csv using CODIGO
3. Find matching HDS file(s) in the HDS folder
4. Extract safety data from PDF/DOCX using text extraction + regex
5. Compare DB entry vs HDS source
6. Output discrepancies to audit_report.txt
"""

import csv
import os
import re
import sys
import unicodedata
from pathlib import Path
from datetime import datetime

# --- PDF extraction ---
try:
    import fitz  # pymupdf
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False
    print("WARNING: pymupdf not installed. PDF files will be skipped.")

# --- DOCX extraction ---
try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("WARNING: python-docx not installed. DOCX files will be skipped.")


# ============================================================
# CONFIGURATION
# ============================================================
BASE_DIR = Path(r"c:\Users\QB_DESARROLLO\Desktop\SGA_dev")
MATRIZ_CSV = BASE_DIR / "data" / "MATRIZ - HDS Y FT (ACTUALIZACION)(Hoja2).csv"
DATABASE_CSV = BASE_DIR / "original_data" / "Unified_GHS_Database.csv"
HDS_FOLDER = Path(r"C:\Users\QB_DESARROLLO\Desktop\Nueva carpeta")
OUTPUT_REPORT = BASE_DIR / "scripts" / "audit_report.txt"


# ============================================================
# TEXT NORMALIZATION UTILITIES
# ============================================================
def normalize_text(text):
    """Normalize text: lowercase, remove accents, strip extra whitespace."""
    if not text:
        return ""
    # Remove accents
    nfkd = unicodedata.normalize('NFKD', text)
    text = ''.join(c for c in nfkd if not unicodedata.combining(c))
    # Lowercase and strip
    text = text.lower().strip()
    # Collapse whitespace
    text = re.sub(r'\s+', ' ', text)
    return text


def normalize_for_matching(name):
    """Normalize a product name for matching against filenames."""
    n = normalize_text(name)
    # Remove common prefixes/suffixes
    n = re.sub(r'\bs/?p\b', '', n)  # remove S/P
    n = re.sub(r'\bqb\b', '', n)
    n = re.sub(r'[^\w\s]', ' ', n)  # remove punctuation
    n = re.sub(r'\s+', ' ', n).strip()
    return n


def normalize_filename(filepath):
    """Extract and normalize the product name from an HDS filename."""
    name = Path(filepath).stem
    # Remove common HDS prefixes
    name = re.sub(r'^HDS\s*[-–—]\s*', '', name, flags=re.IGNORECASE)
    name = re.sub(r'^HDS\s+', '', name, flags=re.IGNORECASE)
    name = re.sub(r'^FORMATO\s+HDS\s+', '', name, flags=re.IGNORECASE)
    # Remove trailing numbers that are version indicators (e.g., "1", " A")
    name = re.sub(r'\s+\d+$', '', name)
    name = re.sub(r'\s+[Aa]$', '', name)
    return normalize_for_matching(name)


# ============================================================
# TEXT EXTRACTION FROM FILES
# ============================================================
def extract_text_from_pdf(filepath):
    """Extract all text from a PDF file using pymupdf."""
    if not HAS_PYMUPDF:
        return ""
    try:
        doc = fitz.open(str(filepath))
        text_parts = []
        for page in doc:
            text_parts.append(page.get_text())
        doc.close()
        return '\n'.join(text_parts)
    except Exception as e:
        return f"[ERROR extracting PDF: {e}]"


def extract_text_from_docx(filepath):
    """Extract all text from a DOCX file using python-docx."""
    if not HAS_DOCX:
        return ""
    try:
        doc = docx.Document(str(filepath))
        text_parts = []
        for para in doc.paragraphs:
            text_parts.append(para.text)
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_parts.append(cell.text)
        return '\n'.join(text_parts)
    except Exception as e:
        return f"[ERROR extracting DOCX: {e}]"


def extract_text(filepath):
    """Extract text from a file based on its extension."""
    ext = Path(filepath).suffix.lower()
    if ext == '.pdf':
        return extract_text_from_pdf(filepath)
    elif ext == '.docx':
        return extract_text_from_docx(filepath)
    return ""


# ============================================================
# SAFETY DATA EXTRACTION FROM TEXT
# ============================================================
def extract_h_statements(text):
    """Extract H-statement codes from text (e.g., H226, H302, H315)."""
    # Match H followed by 3 digits
    matches = re.findall(r'\bH\s*(\d{3})\b', text, re.IGNORECASE)
    # Normalize to H### format
    h_codes = sorted(set(f"H{m}" for m in matches))
    return h_codes


def extract_p_statements(text):
    """Extract P-statement codes from text (e.g., P210, P261, P280)."""
    matches = re.findall(r'\bP\s*(\d{3})\b', text, re.IGNORECASE)
    p_codes = sorted(set(f"P{m}" for m in matches))
    return p_codes


def extract_signal_word(text):
    """Extract the signal word from text."""
    text_upper = text.upper()
    # Check for PELIGRO / DANGER first (higher severity)
    if re.search(r'\bPELIGRO\b', text_upper):
        return "PELIGRO"
    if re.search(r'\bDANGER\b', text_upper):
        return "PELIGRO"
    # Then ATENCION / WARNING
    if re.search(r'\bATENCI[OÓ]N\b', text_upper):
        return "ATENCION"
    if re.search(r'\bWARNING\b', text_upper):
        return "ATENCION"
    return "No Aplicable"


def extract_cas_number(text):
    """Extract CAS numbers from text."""
    # CAS format: 2-7 digits, dash, 2 digits, dash, 1 digit
    matches = re.findall(r'\b(\d{2,7}-\d{2}-\d)\b', text)
    return sorted(set(matches))


def extract_safety_data(text):
    """Extract all safety data from text."""
    return {
        'signal_word': extract_signal_word(text),
        'h_statements': extract_h_statements(text),
        'p_statements': extract_p_statements(text),
        'cas_numbers': extract_cas_number(text),
    }


# ============================================================
# DATABASE LOADING
# ============================================================
def load_matriz(filepath):
    """Load the MATRIZ CSV and return products where ETIQUETA = REVISADO."""
    products = []
    try:
        with open(filepath, 'r', encoding='utf-8-sig') as f:
            reader = csv.DictReader(f)
            for row in reader:
                etiqueta = (row.get('ETIQUETA') or '').strip()
                if etiqueta == 'REVISADO':
                    products.append({
                        'codigo': (row.get('CODIGO') or '').strip(),
                        'nombre': (row.get('NOMBRE') or '').strip(),
                        'hds_status': (row.get('HDS') or '').strip(),
                        'fecha': (row.get('FECHA HDS') or '').strip(),
                    })
    except Exception as e:
        print(f"Error loading MATRIZ CSV: {e}")
        # Try with latin-1
        try:
            with open(filepath, 'r', encoding='latin-1') as f:
                reader = csv.DictReader(f)
                for row in reader:
                    etiqueta = (row.get('ETIQUETA') or '').strip()
                    if etiqueta == 'REVISADO':
                        products.append({
                            'codigo': (row.get('CODIGO') or '').strip(),
                            'nombre': (row.get('NOMBRE') or '').strip(),
                            'hds_status': (row.get('HDS') or '').strip(),
                            'fecha': (row.get('FECHA HDS') or '').strip(),
                        })
        except Exception as e2:
            print(f"Error loading MATRIZ CSV with latin-1: {e2}")
    return products


def load_database(filepath):
    """Load the Unified GHS Database CSV into a dict keyed by Codigo interno."""
    db = {}
    encodings_to_try = ['utf-8-sig', 'utf-8', 'latin-1', 'cp1252']
    
    for enc in encodings_to_try:
        try:
            with open(filepath, 'r', encoding=enc) as f:
                content = f.read()
            break
        except (UnicodeDecodeError, UnicodeError):
            continue
    else:
        print(f"ERROR: Could not read database file with any encoding")
        return db

    # The CSV has multiline fields, so we need careful parsing
    reader = csv.reader(content.splitlines())
    headers = None
    current_row = None
    
    for line_parts in reader:
        if headers is None:
            headers = [h.strip() for h in line_parts]
            continue
        
        # Check if this is a new row (has a codigo in first column) or continuation
        if len(line_parts) >= len(headers):
            if current_row:
                _store_db_row(db, headers, current_row)
            current_row = line_parts
        elif current_row:
            # Continuation line - append to previous row's fields
            # Merge into the last fields
            for i, part in enumerate(line_parts):
                if i < len(current_row):
                    current_row[i] += ' ' + part
                else:
                    current_row.append(part)
        else:
            current_row = line_parts
    
    if current_row:
        _store_db_row(db, headers, current_row)
    
    return db


def _store_db_row(db, headers, row):
    """Store a database row into the db dict."""
    record = {}
    for i, h in enumerate(headers):
        if i < len(row):
            record[h] = row[i].strip()
        else:
            record[h] = ''
    
    codigo = record.get('Codigo interno', '').strip()
    if not codigo:
        return
    
    # Parse H-statements from the record
    h_raw = record.get('H-Statements', '')
    h_codes = extract_h_statements(h_raw)
    
    # Parse signal word
    signal = record.get('Signal Word', '').strip()
    
    # Parse CAS
    cas = record.get('CAS Number', '').strip()
    
    # Pictogram columns
    pictograms = []
    picto_cols = [
        'Exclamación', 'Corrosión', 'Peligro para la salud',
        'Calavera', 'Llama', 'Ambiente',
        'Llama sobre círculo', 'Bomba explotando', 'Cilindro de gas'
    ]
    for col in picto_cols:
        # Try with and without suffix
        val = record.get(col, '')
        if not val:
            # Try with (PICTOGRAMA ) suffix
            val = record.get(f'{col} (PICTOGRAMA )', '')
        if val and val.strip().upper() in ('X', '1'):
            pictograms.append(col)
    
    entry = {
        'codigo': codigo,
        'nombre': record.get('Chemical Name', ''),
        'cas': cas,
        'signal_word': signal,
        'h_statements': h_codes,
        'h_raw': h_raw,
        'pictograms': pictograms,
        'source': record.get('Source_File', ''),
    }
    
    # Store - a codigo may appear multiple times (Master vs Warehouse)
    if codigo not in db:
        db[codigo] = entry
    else:
        # Prefer Master source
        if entry.get('source', '') == 'Master':
            db[codigo] = entry


# ============================================================
# HDS FILE INDEX
# ============================================================
def build_hds_index(hds_folder):
    """Build an index of all HDS files: {normalized_name: [filepath, ...]}."""
    index = {}
    file_list = []
    
    for root, dirs, files in os.walk(hds_folder):
        for f in files:
            ext = Path(f).suffix.lower()
            if ext in ('.pdf', '.docx'):
                full_path = os.path.join(root, f)
                norm_name = normalize_filename(f)
                file_list.append((norm_name, full_path, f))
                
                if norm_name not in index:
                    index[norm_name] = []
                index[norm_name].append(full_path)
    
    return index, file_list


def find_hds_file(product_name, hds_index, hds_file_list):
    """Find the best matching HDS file for a product name."""
    norm_product = normalize_for_matching(product_name)
    
    if not norm_product:
        return None
    
    # 1. Exact match
    if norm_product in hds_index:
        return hds_index[norm_product][0]
    
    # 2. Check if product name is contained in filename or vice versa
    best_match = None
    best_score = 0
    
    for norm_name, filepath, original_name in hds_file_list:
        if not norm_name:
            continue
        
        # Check containment both ways
        if norm_product in norm_name or norm_name in norm_product:
            # Score by length similarity
            score = min(len(norm_product), len(norm_name)) / max(len(norm_product), len(norm_name))
            if score > best_score:
                best_score = score
                best_match = filepath
        
        # Also check individual significant words
        product_words = set(norm_product.split())
        file_words = set(norm_name.split())
        
        # Remove common short words
        stop_words = {'de', 'la', 'el', 'en', 'y', 'a', 'con', 'para', 'por', 'del', 'las', 'los', 'un', 'una'}
        product_words -= stop_words
        file_words -= stop_words
        
        if product_words and file_words:
            common = product_words & file_words
            if common:
                # Jaccard-like score
                score = len(common) / max(len(product_words), len(file_words))
                if score > best_score and score >= 0.5:
                    best_score = score
                    best_match = filepath
    
    if best_score >= 0.4:
        return best_match
    
    return None


# ============================================================
# COMPARISON LOGIC
# ============================================================
def compare_data(db_entry, hds_data, hds_filepath):
    """Compare database entry vs HDS extracted data. Returns list of issues."""
    issues = []
    
    # --- Signal Word ---
    db_signal = normalize_text(db_entry.get('signal_word', ''))
    hds_signal_raw = hds_data.get('signal_word', '')
    hds_signal = normalize_text(hds_signal_raw)
    
    # Normalize signal word variants
    db_signal_norm = db_signal.replace('atencion', 'atencion').replace('no aplicable', 'no aplicable')
    hds_signal_norm = hds_signal.replace('atencion', 'atencion').replace('no aplicable', 'no aplicable')
    
    if db_signal_norm and hds_signal_norm and db_signal_norm != 'no aplicable':
        if db_signal_norm != hds_signal_norm and hds_signal_norm != 'no aplicable':
            issues.append({
                'field': 'Signal Word',
                'severity': 'HIGH',
                'db_value': db_entry.get('signal_word', ''),
                'hds_value': hds_signal_raw,
                'detail': f"DB says '{db_entry.get('signal_word', '')}' but HDS says '{hds_signal_raw}'"
            })
    
    # --- H-Statements ---
    db_h = set(db_entry.get('h_statements', []))
    hds_h = set(hds_data.get('h_statements', []))
    
    # Only compare if HDS actually has H-statements (text extraction worked)
    if hds_h:
        # H-codes in DB but NOT in HDS (DB has extra)
        extra_in_db = db_h - hds_h
        if extra_in_db:
            issues.append({
                'field': 'H-Statements',
                'severity': 'MEDIUM',
                'db_value': ', '.join(sorted(extra_in_db)),
                'hds_value': '(not found in HDS)',
                'detail': f"H-codes in DB but NOT found in HDS: {', '.join(sorted(extra_in_db))}"
            })
        
        # H-codes in HDS but NOT in DB (DB is missing)
        missing_in_db = hds_h - db_h
        if missing_in_db:
            issues.append({
                'field': 'H-Statements',
                'severity': 'HIGH',
                'db_value': '(missing)',
                'hds_value': ', '.join(sorted(missing_in_db)),
                'detail': f"H-codes in HDS but MISSING from DB: {', '.join(sorted(missing_in_db))}"
            })
    
    # --- CAS Number ---
    db_cas = db_entry.get('cas', '').strip()
    hds_cas_list = hds_data.get('cas_numbers', [])
    
    if db_cas and db_cas.lower() not in ('mezcla', 'no aplicable', 'no aplica', '-', 'sustancia', ''):
        # Check if DB CAS appears in HDS CAS list
        db_cas_numbers = re.findall(r'\d{2,7}-\d{2}-\d', db_cas)
        if db_cas_numbers and hds_cas_list:
            for db_c in db_cas_numbers:
                if db_c not in hds_cas_list:
                    issues.append({
                        'field': 'CAS Number',
                        'severity': 'MEDIUM',
                        'db_value': db_c,
                        'hds_value': ', '.join(hds_cas_list) if hds_cas_list else '(none found)',
                        'detail': f"DB CAS '{db_c}' not found in HDS CAS numbers: {', '.join(hds_cas_list)}"
                    })
    
    return issues


# ============================================================
# MAIN AUDIT
# ============================================================
def run_audit():
    """Run the full HDS audit."""
    print("=" * 70)
    print("  HDS DATABASE AUDIT")
    print(f"  Started: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print("=" * 70)
    
    # 1. Load MATRIZ
    print("\n[1/5] Loading MATRIZ CSV...")
    products = load_matriz(MATRIZ_CSV)
    print(f"  → Found {len(products)} products with ETIQUETA = REVISADO")
    
    if not products:
        print("ERROR: No products found. Check CSV path and encoding.")
        return
    
    # 2. Load Database
    print("\n[2/5] Loading SGA Database...")
    db = load_database(DATABASE_CSV)
    print(f"  → Loaded {len(db)} unique product codes")
    
    # 3. Build HDS file index
    print("\n[3/5] Indexing HDS files...")
    hds_index, hds_file_list = build_hds_index(HDS_FOLDER)
    print(f"  → Indexed {len(hds_file_list)} files ({len(hds_index)} unique names)")
    
    # 4. Audit each product
    print("\n[4/5] Auditing products...")
    
    results = []
    stats = {
        'total': len(products),
        'audited': 0,
        'not_in_db': 0,
        'no_hds_file': 0,
        'extraction_error': 0,
        'with_issues': 0,
        'clean': 0,
        'total_issues': 0,
    }
    
    for i, product in enumerate(products):
        codigo = product['codigo']
        nombre = product['nombre']
        
        if (i + 1) % 25 == 0 or i == 0:
            print(f"  Processing {i+1}/{len(products)}: {nombre[:40]}...")
        
        result = {
            'codigo': codigo,
            'nombre': nombre,
            'fecha': product['fecha'],
            'status': 'OK',
            'issues': [],
            'hds_file': None,
            'notes': [],
        }
        
        # Check if product exists in database
        if codigo not in db:
            result['status'] = 'NOT_IN_DB'
            result['notes'].append(f"Product code '{codigo}' not found in SGA database")
            stats['not_in_db'] += 1
            results.append(result)
            continue
        
        db_entry = db[codigo]
        
        # Check if DB entry has meaningful data (skip if all "No Aplicable" or "-")
        db_signal = (db_entry.get('signal_word') or '').strip()
        db_h = db_entry.get('h_statements', [])
        
        if db_signal in ('-', '') and not db_h:
            result['notes'].append("DB entry has no safety data (may be pending)")
        
        # Find matching HDS file
        hds_file = find_hds_file(nombre, hds_index, hds_file_list)
        
        if not hds_file:
            # Also try with the DB chemical name
            db_name = db_entry.get('nombre', '')
            if db_name:
                hds_file = find_hds_file(db_name, hds_index, hds_file_list)
        
        if not hds_file:
            result['status'] = 'NO_HDS_FILE'
            result['notes'].append(f"No matching HDS file found for '{nombre}'")
            stats['no_hds_file'] += 1
            results.append(result)
            continue
        
        result['hds_file'] = hds_file
        
        # Extract text from HDS file
        text = extract_text(hds_file)
        
        if not text or text.startswith('[ERROR'):
            result['status'] = 'EXTRACTION_ERROR'
            result['notes'].append(f"Could not extract text from: {hds_file}")
            stats['extraction_error'] += 1
            results.append(result)
            continue
        
        # Extract safety data from HDS text
        hds_data = extract_safety_data(text)
        
        # Compare DB vs HDS
        issues = compare_data(db_entry, hds_data, hds_file)
        
        stats['audited'] += 1
        
        if issues:
            result['status'] = 'ISSUES_FOUND'
            result['issues'] = issues
            stats['with_issues'] += 1
            stats['total_issues'] += len(issues)
        else:
            result['status'] = 'OK'
            stats['clean'] += 1
        
        results.append(result)
    
    # 5. Generate report
    print("\n[5/5] Generating audit report...")
    generate_report(results, stats)
    print(f"\n{'=' * 70}")
    print(f"  AUDIT COMPLETE")
    print(f"  Report saved to: {OUTPUT_REPORT}")
    print(f"{'=' * 70}")
    print(f"\n  Summary:")
    print(f"    Total REVISADO products:    {stats['total']}")
    print(f"    Successfully audited:       {stats['audited']}")
    print(f"    Products with issues:       {stats['with_issues']}")
    print(f"    Clean (no issues):          {stats['clean']}")
    print(f"    Not in database:            {stats['not_in_db']}")
    print(f"    No HDS file found:          {stats['no_hds_file']}")
    print(f"    Extraction errors:          {stats['extraction_error']}")
    print(f"    Total issues found:         {stats['total_issues']}")


def generate_report(results, stats):
    """Generate the audit report text file."""
    with open(OUTPUT_REPORT, 'w', encoding='utf-8') as f:
        f.write("=" * 80 + "\n")
        f.write("  HDS DATABASE AUDIT REPORT\n")
        f.write(f"  Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        f.write("=" * 80 + "\n\n")
        
        # Summary
        f.write("SUMMARY\n")
        f.write("-" * 40 + "\n")
        f.write(f"Total REVISADO products:    {stats['total']}\n")
        f.write(f"Successfully audited:       {stats['audited']}\n")
        f.write(f"Products with issues:       {stats['with_issues']}\n")
        f.write(f"Clean (no issues):          {stats['clean']}\n")
        f.write(f"Not in database:            {stats['not_in_db']}\n")
        f.write(f"No HDS file found:          {stats['no_hds_file']}\n")
        f.write(f"Extraction errors:          {stats['extraction_error']}\n")
        f.write(f"Total issues found:         {stats['total_issues']}\n\n")
        
        # === SECTION 1: Products with Issues ===
        issues_results = [r for r in results if r['status'] == 'ISSUES_FOUND']
        if issues_results:
            f.write("\n" + "=" * 80 + "\n")
            f.write("  SECTION 1: PRODUCTS WITH DISCREPANCIES (NEED MANUAL REVIEW)\n")
            f.write("=" * 80 + "\n\n")
            
            for r in issues_results:
                f.write(f"PRODUCT: {r['nombre']}\n")
                f.write(f"  Code:     {r['codigo']}\n")
                f.write(f"  HDS File: {r['hds_file']}\n")
                f.write(f"  Fecha:    {r['fecha']}\n")
                
                for issue in r['issues']:
                    severity_icon = "⚠️" if issue['severity'] == 'HIGH' else "ℹ️"
                    f.write(f"  [{issue['severity']}] {issue['field']}:\n")
                    f.write(f"    → {issue['detail']}\n")
                
                f.write("\n" + "-" * 60 + "\n\n")
        
        # === SECTION 2: Products NOT in Database ===
        not_in_db = [r for r in results if r['status'] == 'NOT_IN_DB']
        if not_in_db:
            f.write("\n" + "=" * 80 + "\n")
            f.write("  SECTION 2: PRODUCTS NOT FOUND IN DATABASE\n")
            f.write("=" * 80 + "\n\n")
            
            for r in not_in_db:
                f.write(f"  {r['codigo']:20s} {r['nombre']}\n")
            f.write("\n")
        
        # === SECTION 3: No HDS File Found ===
        no_hds = [r for r in results if r['status'] == 'NO_HDS_FILE']
        if no_hds:
            f.write("\n" + "=" * 80 + "\n")
            f.write("  SECTION 3: NO MATCHING HDS FILE FOUND\n")
            f.write("=" * 80 + "\n\n")
            
            for r in no_hds:
                f.write(f"  {r['codigo']:20s} {r['nombre']}\n")
            f.write("\n")
        
        # === SECTION 4: Extraction Errors ===
        errors = [r for r in results if r['status'] == 'EXTRACTION_ERROR']
        if errors:
            f.write("\n" + "=" * 80 + "\n")
            f.write("  SECTION 4: FILE EXTRACTION ERRORS\n")
            f.write("=" * 80 + "\n\n")
            
            for r in errors:
                f.write(f"  {r['codigo']:20s} {r['nombre']}\n")
                f.write(f"    File: {r['hds_file']}\n")
                for note in r['notes']:
                    f.write(f"    Note: {note}\n")
            f.write("\n")
        
        # === SECTION 5: Clean Products ===
        clean = [r for r in results if r['status'] == 'OK']
        if clean:
            f.write("\n" + "=" * 80 + "\n")
            f.write("  SECTION 5: CLEAN PRODUCTS (NO ISSUES FOUND)\n")
            f.write("=" * 80 + "\n\n")
            
            for r in clean:
                f.write(f"  {r['codigo']:20s} {r['nombre']}\n")
            f.write("\n")
        
        f.write("\n" + "=" * 80 + "\n")
        f.write("  END OF REPORT\n")
        f.write("=" * 80 + "\n")


if __name__ == '__main__':
    run_audit()
